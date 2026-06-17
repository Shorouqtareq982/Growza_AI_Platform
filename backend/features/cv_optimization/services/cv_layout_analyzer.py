import io
import re
from typing import Dict, List, Optional, Set, Tuple, TypedDict

from fastapi import UploadFile
import numpy as np
import pymupdf
import math
from statistics import mean
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.text.run import Run
from sklearn.cluster import KMeans
from ..helpers.multi_column import column_boxes, have_columns as have_columns_func

from shared.helpers.file_validation import FileValidator
from ..schemas.layout_analysis_schema import CVLayoutAnalysis, PageMargin, PageSize


# ==================== Type Definitions ====================

class PdfPageAnalysisResult(TypedDict):
    """Result from analyzing a single PDF page."""
    have_images: bool
    have_tables: bool
    have_graphics: bool
    have_columns: bool
    information_in_header: bool
    information_in_footer: bool
    page_size: Dict[str, float]
    page_margin: Dict[str, float]


class HeaderFooterResult(TypedDict):
    """Result from header/footer analysis."""
    has_header_info: bool
    has_footer_info: bool
    have_links_in_header: bool


class TextContentResult(TypedDict):
    """Result from text content analysis."""
    extracted_text: str
    fonts_names: List[str]
    font_sizes: List[float]
    have_hyperlinks: bool


class CVLayoutAnalyzer:
    """Analyzes CV layout features from PDF and DOCX files.
    
    Detects layout properties including:
    - Images, tables, graphics, columns
    - Font information and sizes
    - Page dimensions and margins
    - Header and footer content
    - Hyperlinks and document structure
    """

    # ==================== Constants ====================
    # PDF Analysis Constants
    KMEANS_N_CLUSTERS = 2
    KMEANS_INIT = 10
    COLUMN_SEPARATION_THRESHOLD = 200
    MIN_X_POSITIONS_FOR_CLUSTERING = 30
    
    # Page Margin Constants (as percentage of page height)
    HEADER_THRESHOLD_PERCENT = 0.05
    FOOTER_THRESHOLD_PERCENT = 0.95
    
    # Conversion Factors
    POINTS_PER_INCH = 72
    EMU_PER_POINT = 12700
    EMU_PER_INCH = 914400
    
    # XML Namespace URIs
    WPML_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    WPS_NS = "{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}wsp"
    VML_NS = "{urn:schemas-microsoft-com:vml}shape"
    VML_TEXTBOX_NS = "{urn:schemas-microsoft-com:vml}textbox"
    WP_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    RELATIONSHIPS_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
    
    # Validation Constants
    MAX_FILENAME_LENGTH = 100

    # ==================== Main Entry Point ====================
    @staticmethod
    def analyze_file_layout(
        file_buf: io.BytesIO,
        file_name: str,
        file_size_kb: float,
        file_type: str,
    ) -> Dict | Tuple[Dict, str]:
        """Analyze layout of a CV file (PDF or DOCX).
        
        Args:
            file_buf: BytesIO object containing file data
            file_name: Original filename for validation
            file_size_kb: File size in kilobytes
            file_type: MIME type of the file
        
        Returns:
            Dictionary containing layout analysis results for PDF, or tuple(Dictionary, text) for DOCX
            Returns default analysis with None values on error
        """
        try:
            name_lower = (file_name or "").lower()
            if name_lower.endswith(".pdf"):
                doc = pymupdf.open(stream=file_buf.read(), filetype="pdf")
                try:
                    return CVLayoutAnalyzer.analyze_pdf_layout(doc, file_size_kb, file_type, file_name)
                finally:
                    doc.close()
            elif name_lower.endswith(".docx"):
                doc = Document(file_buf)
                return CVLayoutAnalyzer.analyze_docx_layout(doc, file_size_kb, file_type, file_name)
            else:
                raise ValueError("Unsupported file type for layout analysis")
        except Exception as e:
            # Log the error and return a default analysis indicating failure
            print(f"Error analyzing file layout: {e}")
            return {
                "have_images": None,
                "have_tables": None,
                "have_columns": None,
                "have_graphics": None,
                "have_textboxes": None,
                "information_in_header": None,
                "information_in_footer": None,
                "fonts_used": None,
                "avg_font_size": None,
                "file_size_kb": file_size_kb,
                "file_type": file_type,
                "num_of_pages": None,
                "num_of_sections": None,
                "word_count": None,
                "valid_cv_filename": FileValidator.validate_filename(file_name)[0] if file_name else None,
                "valid_cv_filename_length": len(file_name) <= CVLayoutAnalyzer.MAX_FILENAME_LENGTH if file_name else None,
                "original_filename": file_name,
                "page_sizes_in_points": None,
                "page_margins_in_inches": None
            }, None

    # ==================== PDF Analysis Methods ====================

    @staticmethod
    def analyze_pdf_layout(
        doc, 
        file_size_kb: float, 
        file_type: str, 
        org_filename: str
    ) -> Tuple[CVLayoutAnalysis, str]:
        """Analyze PDF layout features using multi_column.py for better text extraction.
        
        Args:
            doc: PyMuPDF document object
            file_size_kb: File size in kilobytes
            file_type: Content type of the file
            org_filename: Original file name
        
        Returns:
            Tuple of (CVLayoutAnalysis object, extracted text string)
        """
        have_images = False
        have_tables = False
        have_graphics = False
        have_columns = False
        num_of_doc_words = 0
        information_in_header = False
        information_in_footer = False
        page_margins = []
        page_texts = []
        page_sizes = []
        avg_font_size = 0
        all_font_sizes = []
        all_text_content = []
        valid_filename, _ = FileValidator.validate_filename(org_filename)
        valid_filename_length = len(org_filename) <= CVLayoutAnalyzer.MAX_FILENAME_LENGTH
        num_of_pages = doc.page_count
        fonts_used = set()

        # Analyze each page
        for page in doc:
            page_analysis = CVLayoutAnalyzer._analyze_pdf_page_with_column_boxes(
                page, all_font_sizes, all_text_content, fonts_used
            )
            have_images |= page_analysis["have_images"]
            have_tables |= page_analysis["have_tables"]
            have_graphics |= page_analysis["have_graphics"]
            have_columns |= page_analysis["have_columns"]
            information_in_header |= page_analysis["information_in_header"]
            information_in_footer |= page_analysis["information_in_footer"]
            page_margins.append(page_analysis["page_margin"])
            page_sizes.append(page_analysis["page_size"])
            page_texts.append(page_analysis["page_text"])

        if all_font_sizes:
            avg_font_size = round(sum(all_font_sizes) / len(all_font_sizes), 2)

        # Calculate total word count
        num_of_doc_words = sum(len(text.split()) for text in all_text_content)

        # Convert to Pydantic objects
        def _safe_page_size(ps):
            if not ps:
                return PageSize(width=8.5, height=11.0)
            try:
                return PageSize(**ps)
            except Exception:
                return PageSize(width=8.5, height=11.0)

        def _safe_page_margin(pm):
            defaults = {"top": 1.0, "bottom": 1.0, "left": 1.0, "right": 1.0}
            if not pm:
                return PageMargin(**defaults)
            try:
                merged = {**defaults, **pm}
                return PageMargin(**merged)
            except Exception:
                return PageMargin(**defaults)

        page_sizes_objs = [_safe_page_size(ps) for ps in page_sizes]
        page_margins_objs = [_safe_page_margin(pm) for pm in page_margins]

        extracted_text = CVLayoutAnalyzer._clean_text(page_texts)

        return CVLayoutAnalysis(
            have_images=have_images,
            have_tables=have_tables,
            have_columns=have_columns,
            have_graphics=have_graphics,
            information_in_header=information_in_header,
            information_in_footer=information_in_footer,
            fonts_used=list(fonts_used),
            avg_font_size=avg_font_size,
            file_size_kb=file_size_kb,
            file_type=file_type,
            num_of_pages=num_of_pages,
            num_of_sections=None,
            word_count=num_of_doc_words,
            valid_cv_filename=valid_filename,
            valid_cv_filename_length=valid_filename_length,
            original_filename=org_filename,
            page_sizes_in_points=page_sizes_objs,
            page_margins_in_inches=page_margins_objs,
        ), extracted_text

    @staticmethod
    def _analyze_pdf_page_with_column_boxes(
        page, 
        all_font_sizes: List[float], 
        all_text_content: List[str], 
        fonts_used: Set[str]
    ) -> PdfPageAnalysisResult:
        """Analyze a single PDF page using column_boxes for improved text extraction.
        
        This method uses the multi_column.py column_boxes function to:
        1. Detect multi-column layouts accurately
        2. Extract text in proper reading order
        3. Handle tables separately from body text
        """
        page_rect = page.rect
        page_width = page_rect.width
        page_height = page_rect.height
        
        # Header / footer thresholds
        header_threshold = page_height * CVLayoutAnalyzer.HEADER_THRESHOLD_PERCENT
        footer_threshold = page_height * CVLayoutAnalyzer.FOOTER_THRESHOLD_PERCENT

        have_images = bool(page.get_images())
        table_objects = page.find_tables()
        have_tables = len(table_objects.tables) > 0
        table_bboxes = [table.bbox for table in table_objects.tables] if have_tables else []
        have_graphics = len(page.get_drawings()) > 0
        information_in_header = False
        information_in_footer = False
        page_margin = {}
        page_text = ""

        # Extract fonts
        fonts = page.get_fonts()
        for font in fonts:
            raw_font_name = font[3]
            clean_name = re.sub(r"^[A-Z]{6}\+", "", raw_font_name)
            # Remove common PDF postfixes
            clean_name = re.sub(r"(PSMT|PS|MT)$", "", clean_name, flags=re.I)
            # Remove style descriptors
            clean_name = re.sub(
                r"[-_,]?(?:"
                r"Bold|Italic|Regular|Oblique|Roman|Medium|"
                r"Light|SemiBold|DemiBold|ExtraBold|Black|"
                r"Thin|Condensed|Narrow|Book|Heavy|"
                r"BoldItalic|LightItalic|MediumItalic"
                r")$",
                "",
                clean_name,
                flags=re.I,
            )
            # Replace separators with spaces
            clean_name = re.sub(r"[-_,]+", " ", clean_name)
            # Collapse multiple spaces
            clean_name = re.sub(r"\s+", " ", clean_name)
            
            fonts_used.add(clean_name)

        # ============================================================
        # KEY IMPROVEMENT: Use column_boxes from multi_column.py
        # ============================================================
        # This provides intelligent multi-column detection and proper text ordering
        try:
            text_bboxes = column_boxes(
                page,
                footer_margin=footer_threshold,
                header_margin=header_threshold,
                no_image_text=True,
                avoid=table_bboxes  # Avoid extracting text from table areas
            )
            # If we got multiple bboxes, we have columns
            # have_columns = len(text_bboxes) > 1 and CVLayoutAnalyzer._has_column_layout(text_bboxes, page_width)
        except Exception as e:
            # Fallback to basic extraction if column_boxes fails
            print(f"Warning: column_boxes failed: {e}")
            text_bboxes = []
            # have_columns = False

        # Extract text from the ordered bounding boxes
        page_lines = []
        raw_blocks = page.get_text("rawdict")["blocks"]
        text_blocks = [b for b in raw_blocks if b["type"] == 0]

        def _normalize_line(line_text: str) -> str:
            line_text = line_text.strip()
            if line_text.startswith(("•", "-", "▪", "●", "◦")):
                return f"- {line_text.lstrip('•-▪●◦').strip()}"
            return line_text

        def _extract_block_lines(block: Dict) -> List[str]:
            lines = []
            for line in block.get("lines", []):
                line_text = ""
                for span in line.get("spans", []):
                    chars = span.get("chars", [])
                    text = "".join(ch.get("c", "") for ch in chars).strip()
                    if text:
                        line_text += text + " "

                line_text = _normalize_line(line_text)
                if line_text:
                    lines.append(line_text)

            return lines

        if have_tables:
            table_segments = []
            for table in table_objects.tables:
                table_lines = CVLayoutAnalyzer._extract_table_content(page, table)
                if table_lines:
                    table_segments.append((table.bbox, table_lines))

            non_table_blocks = CVLayoutAnalyzer._filter_blocks_in_tables(text_blocks, table_bboxes)
            text_segments = [(block["bbox"], _extract_block_lines(block)) for block in non_table_blocks]

            ordered_segments = sorted(
                [segment for segment in text_segments + table_segments if segment[1]],
                key=lambda item: (item[0][1], item[0][0])
            )

            for bbox, lines in ordered_segments:
                if bbox[1] < header_threshold:
                    information_in_header = True
                if bbox[3] > footer_threshold:
                    information_in_footer = True

                for line in lines:
                    page_lines.append(line)
                    all_text_content.append(line)
        else:
            # Fallback: use standard block extraction if column_boxes didn't work
            if text_bboxes:
                for bbox in text_bboxes:
                    # Extract text from this bbox
                    text = page.get_text("text", clip=bbox, sort=True)
                    if text.strip():
                        lines = text.strip().split('\n')
                        # Normalize bullets
                        for line in lines:
                            line = _normalize_line(line)
                            if line:
                                page_lines.append(line)
                                all_text_content.append(line)
                        
                    # Check if this bbox is in header/footer area
                    if bbox.y0 < header_threshold:
                        information_in_header = True
                    if bbox.y1 > footer_threshold:
                        information_in_footer = True
            else:
                non_table_blocks = CVLayoutAnalyzer._filter_blocks_in_tables(text_blocks, table_bboxes)
                
                for block in non_table_blocks:
                    block_top = block["bbox"][1]
                    block_bottom = block["bbox"][3]
                    
                    if block_top < header_threshold:
                        information_in_header = True
                    if block_bottom > footer_threshold:
                        information_in_footer = True
                    
                    for line_text in _extract_block_lines(block):
                        page_lines.append(line_text)
                        all_text_content.append(line_text)

        # Collect font sizes from all text blocks
        for block in text_blocks:
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    if span.get("size", 0) > 0:
                        all_font_sizes.append(span["size"])

        # Build final page text
        page_text = "\n".join(page_lines)
        page_text = CVLayoutAnalyzer._attach_links_once(page, page_text)

        # Estimate page margins from text blocks
        if text_blocks:
            # text_blocks = [
            #     b for b in text_blocks
            #     if (b["bbox"][2] - b["bbox"][0]) > 50
            #     and (b["bbox"][3] - b["bbox"][1]) > 10
            # ]
            left_margin = min(b["bbox"][0] for b in text_blocks)
            top_margin = min(b["bbox"][1] for b in text_blocks)
            right_margin = page_width - max(b["bbox"][2] for b in text_blocks)
            bottom_margin = page_height - max(b["bbox"][3] for b in text_blocks)

            page_margin = {
                "left": max(0, round(left_margin / CVLayoutAnalyzer.POINTS_PER_INCH, 2)),
                "top": max(0, round(top_margin / CVLayoutAnalyzer.POINTS_PER_INCH, 2)),
                "right": max(0, round(right_margin / CVLayoutAnalyzer.POINTS_PER_INCH, 2)),
                "bottom": max(0, round(bottom_margin / CVLayoutAnalyzer.POINTS_PER_INCH, 2)),
            }

        have_columns = have_columns_func(page)
        return {
            "have_images": have_images,
            "have_tables": have_tables,
            "have_graphics": have_graphics,
            "have_columns": have_columns,
            "information_in_header": information_in_header,
            "information_in_footer": information_in_footer,
            "page_size": {
                "width": round(page_width, 2),
                "height": round(page_height, 2)
            },
            "page_margin": page_margin,
            "page_text": page_text
        }
    
    @staticmethod
    def _has_column_layout(bboxes: List, page_width: float) -> bool:
        """Determine if bboxes represent a multi-column layout.
        
        Args:
            bboxes: List of text bounding boxes from column_boxes
            page_width: Width of the page
        
        Returns:
            True if layout appears to be multi-column
        """
        if len(bboxes) < 2:
            return False
        
        # Check if we have boxes with significant horizontal separation
        # Get x-coordinates of boxes
        x_positions = [(bbox.x0 + bbox.x1) / 2 for bbox in bboxes]
        
        if len(set(x_positions)) < 2:
            return False
        
        # Check for significant separation (at least 20% of page width)
        min_separation = page_width * 0.2
        sorted_x = sorted(set(x_positions))
        
        for i in range(len(sorted_x) - 1):
            if sorted_x[i + 1] - sorted_x[i] > min_separation:
                return True
        
        return False
    
    @staticmethod
    def _filter_blocks_in_tables(blocks: List[Dict], table_bboxes: List) -> List[Dict]:
        """Filter out text blocks that are inside table boundaries.
        
        Only filters blocks whose Y-coordinates are within table Y-range.
        Blocks that come before or after tables are kept.
        
        Args:
            blocks: List of text blocks from page
            table_bboxes: List of table bounding boxes (x0, y0, x1, y1)
        
        Returns:
            List of blocks that are NOT inside tables
        """
        def is_inside_table(block_bbox: Tuple) -> bool:
            """Check if a block's center Y is within any table's Y range AND X overlaps."""
            block_x0, block_y0, block_x1, block_y1 = block_bbox
            block_y_center = (block_y0 + block_y1) / 2
            
            for table_bbox in table_bboxes:
                table_x0, table_y0, table_x1, table_y1 = table_bbox
                if (table_y0 <= block_y_center <= table_y1 and
                    block_x0 < table_x1 and block_x1 > table_x0):
                    return True
            return False
        
        return [b for b in blocks if not is_inside_table(b["bbox"])]
 
    @staticmethod
    def _clean_table_cell(text: Optional[str]) -> str:
        """Normalize a single table cell value."""
        if text is None:
            return ""

        text = text.replace("\n", " ")
        text = re.sub(r"\s+", " ", text)

        return text.strip()

    @staticmethod
    def _remove_empty_columns(table: List[List[str]]) -> List[List[str]]:
        """Remove columns that contain only empty cells."""
        if not table:
            return table

        cols_to_keep = []
        num_cols = max(len(row) for row in table)

        for col_idx in range(num_cols):
            column_values = []

            for row in table:
                if col_idx < len(row):
                    column_values.append(row[col_idx].strip())

            if any(value != "" for value in column_values):
                cols_to_keep.append(col_idx)

        cleaned = []

        for row in table:
            cleaned.append([
                row[i] if i < len(row) else ""
                for i in cols_to_keep
            ])

        return cleaned

    @staticmethod
    def _extract_table_content(page, table) -> List[str]:
        """Extract text from a single table with structure preservation."""
        table_lines = []
        section_titles = []
        
        try:
            rows = []

            for row in table.rows:
                values = []

                for cell in row.cells:
                    if cell is None:
                        values.append("")
                        continue

                    rect = pymupdf.Rect(cell)
                    text = page.get_text("text", clip=rect)
                    values.append(CVLayoutAnalyzer._clean_table_cell(text))

                non_empty = [value for value in values if value.strip()]

                if len(non_empty) == 1 and non_empty[0].isupper():
                    section_titles.append(non_empty[0])
                    continue

                rows.append(values)

            rows = CVLayoutAnalyzer._remove_empty_columns(rows)

            if not rows and not section_titles:
                return table_lines
                        
            for row_idx, row in enumerate(rows):
                row_text = " | ".join(cell for cell in row if cell.strip())
                
                if row_text.strip():
                    table_lines.append(row_text)
            
            for title in section_titles:
                table_lines.append(title)

            table_lines.append("")
            
        except Exception:
            try:
                text = table.get_text().strip()
                if text:
                    table_lines.extend(text.split("\n"))
                    table_lines.append("")
            except Exception:
                pass
        
        return table_lines
 
    @staticmethod
    def _attach_links_once(page, page_text: str) -> str:
        """Attach hyperlink URLs to their anchor text in the page text."""
        links = page.get_links()
        replacements = []
        seen_uris = set()
 
        for link in links:
            uri = link.get("uri")
            rect = link.get("from")
 
            if not uri or uri in seen_uris or not rect:
                continue
 
            anchor_text = page.get_text("text", clip=rect).strip()
 
            if anchor_text:
                replacements.append((anchor_text, f"{anchor_text} ({uri})"))
                seen_uris.add(uri)
 
        for original, replaced in replacements:
            page_text = page_text.replace(original, replaced, 1)
 
        return page_text
    # ==================== DOCX Analysis Methods ====================

    @staticmethod
    def analyze_docx_layout(
        doc: Document,
        file_size_kb: float,
        file_type: str,
        org_filename: str
    ) -> Tuple[CVLayoutAnalysis, str]:
        """Analyze DOCX layout features and extract text content.
        
        Args:
            doc: python-docx Document object
            file_size_kb: File size in kilobytes
            file_type: MIME type of the file
            org_filename: Original file name
        
        Returns:
            Tuple of (CVLayoutAnalysis object, extracted text string)
        """
        valid_filename, _ = FileValidator.validate_filename(org_filename)
        valid_filename_length = len(org_filename) <= CVLayoutAnalyzer.MAX_FILENAME_LENGTH
        num_of_pages = CVLayoutAnalyzer._get_page_count_docx(doc)
        page_sizes, page_margins, have_columns = CVLayoutAnalyzer._extract_page_properties(doc, num_of_pages)
        content_data = CVLayoutAnalyzer._analyze_text_content(doc)
        header_footer_data = CVLayoutAnalyzer._analyze_headers_footers(doc)

        fonts_used = content_data["fonts_names"]
        all_font_sizes = content_data["font_sizes"]
        avg_font_size = round(sum(all_font_sizes) / len(all_font_sizes), 2) if all_font_sizes else 0
        num_of_doc_words = CVLayoutAnalyzer.word_count(content_data["extracted_text"])

        # Detect document features
        have_graphics = CVLayoutAnalyzer._detect_graphics(doc)
        have_images = CVLayoutAnalyzer._detect_images(doc)  
        have_textboxes = CVLayoutAnalyzer._detect_textboxes(doc)
        have_tables = len(doc.tables) > 0


        return CVLayoutAnalysis(
            have_images=have_images,
            have_tables=have_tables,
            have_columns=have_columns,
            have_graphics=have_graphics,
            have_textboxes=have_textboxes,
            information_in_header=header_footer_data["has_header_info"],
            information_in_footer=header_footer_data["has_footer_info"],
            fonts_used=set(fonts_used),
            avg_font_size=avg_font_size,
            file_size_kb=file_size_kb,
            file_type=file_type,
            num_of_pages=num_of_pages,
            num_of_sections=len(doc.sections),
            word_count=num_of_doc_words,
            valid_cv_filename=valid_filename,
            valid_cv_filename_length=valid_filename_length,
            original_filename=org_filename,
            page_sizes_in_points=page_sizes,
            page_margins_in_inches=page_margins
        ), content_data["extracted_text"]
    
    @staticmethod
    def _extract_page_properties(doc: Document, num_pages: int) -> Tuple[List[PageSize], List[PageMargin], bool]:
        """Extract page sizes, margins, and column information from DOCX sections.
        
        Distributes section properties across all pages proportionally.
        
        Args:
            doc: python-docx Document object
            num_pages: Total number of pages in document
        
        Returns:
            Tuple of (page_sizes list, page_margins list, have_columns bool)
        """
        page_sizes = []
        page_margins = []
        have_columns = False

        # DOCX page size and margins are stored in section properties.
        # If the document exposes no usable section data, fall back to a standard page setup.
        if not getattr(doc, "sections", None):
            default_size = PageSize(width=8.5, height=11.0)
            default_margin = PageMargin(top=1.0, bottom=1.0, left=1.0, right=1.0)
            page_sizes = [default_size for _ in range(num_pages)]
            page_margins = [default_margin for _ in range(num_pages)]
            return page_sizes, page_margins, have_columns

        # Collect section properties
        section_properties = []
        for section in doc.sections:
            page_size = CVLayoutAnalyzer._get_section_page_size(section)
            margins = CVLayoutAnalyzer._get_section_margins(section)
            section_properties.append((page_size, margins))

            # Check for columns
            sect_pr = section._sectPr
            cols_elem = sect_pr.find(f"{CVLayoutAnalyzer.WPML_NS}cols")
            if cols_elem is not None:
                num_cols = cols_elem.get(f"{CVLayoutAnalyzer.WPML_NS}num")
                if num_cols and int(num_cols) > 1:
                    have_columns = True

        # Distribute section properties across pages
        page_sizes, page_margins = CVLayoutAnalyzer._distribute_section_properties(
            section_properties, num_pages
        )

        return page_sizes, page_margins, have_columns

    @staticmethod
    def _get_section_page_size(section) -> Optional[PageSize]:
        """Extract page size from a document section."""
        if section.page_width and section.page_height:
            return PageSize(
                width=round(section.page_width / CVLayoutAnalyzer.EMU_PER_POINT, 2),
                height=round(section.page_height / CVLayoutAnalyzer.EMU_PER_POINT, 2),
            )
        return None

    @staticmethod
    def _get_section_margins(section) -> PageMargin:
        """Extract margins from a document section.
        
        Converts EMU (English Metric Units) to inches.
        """
        return PageMargin(
            top=round(section.top_margin / CVLayoutAnalyzer.EMU_PER_INCH, 2),
            bottom=round(section.bottom_margin / CVLayoutAnalyzer.EMU_PER_INCH, 2),
            left=round(section.left_margin / CVLayoutAnalyzer.EMU_PER_INCH, 2),
            right=round(section.right_margin / CVLayoutAnalyzer.EMU_PER_INCH, 2)
        )


    @staticmethod
    def _distribute_section_properties(
        section_properties: List[Tuple], 
        num_pages: int
    ) -> Tuple[List[PageSize], List[PageMargin]]:
        """Distribute section properties across pages.
        
        Args:
            section_properties: List of (page_size, margins) tuples
            num_pages: Total number of pages
        
        Returns:
            Tuple of (page_sizes list, page_margins list)
        """
        page_sizes = []
        page_margins = []

        if len(section_properties) == 1:
            # Single section: replicate properties for all pages
            page_size, margins = section_properties[0]
            for _ in range(num_pages):
                if page_size:
                    page_sizes.append(page_size)
                page_margins.append(margins)
        elif len(section_properties) == 0:
            # No section properties: use defaults
            default_size = PageSize(width=8.5, height=11.0)
            default_margin = PageMargin(top=1.0, bottom=1.0, left=1.0, right=1.0)
            for _ in range(num_pages):
                page_sizes.append(default_size)
                page_margins.append(default_margin)
        else:
            # Multiple sections: estimate pages per section
            pages_per_section = max(1, num_pages // len(section_properties))
            remaining_pages = num_pages % len(section_properties)
            page_count = 0

            for section_idx, (page_size, margins) in enumerate(section_properties):
                pages_in_section = pages_per_section + (1 if section_idx < remaining_pages else 0)

                for _ in range(pages_in_section):
                    if page_size:
                        page_sizes.append(page_size)
                    page_margins.append(margins)
                    page_count += 1
                    if page_count >= num_pages:
                        break
                if page_count >= num_pages:
                    break

        # Ensure page_sizes and page_margins have matching lengths
        if len(page_sizes) < len(page_margins):
            last_size = page_sizes[-1] if page_sizes else PageSize(width=8.5, height=11.0)
            page_sizes.extend([last_size] * (len(page_margins) - len(page_sizes)))

        return page_sizes, page_margins

    
    @staticmethod
    def _detect_images(doc: Document) -> bool:
        """Detect if DOCX contains images by checking document relationships."""
        try:
            for rel in doc.part.rels.values():
                if hasattr(rel, "target_ref") and "image" in rel.target_ref:
                    return True
        except Exception:
            pass
        return False
    
    @staticmethod
    def _detect_graphics(doc: Document) -> bool:
        """Detect if DOCX contains graphics/shapes (both modern and legacy VML)."""
        body = doc._element.body
        return bool(body.findall(f".//{CVLayoutAnalyzer.WPS_NS}") or 
                   body.findall(f".//{CVLayoutAnalyzer.VML_NS}"))
    
    @staticmethod
    def _detect_textboxes(doc: Document) -> bool:
        """Detect if DOCX contains textboxes."""
        body = doc._element.body
        return bool(body.findall(f".//{CVLayoutAnalyzer.WP_NS}txbxContent") or
                   body.findall(f".//{CVLayoutAnalyzer.VML_TEXTBOX_NS}"))

    @staticmethod
    def _analyze_headers_footers(doc: Document) -> HeaderFooterResult:
        """Analyze headers and footers for content and hyperlinks.
        
        Returns:
            HeaderFooterResult with header_info, footer_info, have_links_in_header
        """
        header_info = False
        footer_info = False
        have_links_in_header = False

        for section in doc.sections:
            # Check headers
            header_info, has_links = CVLayoutAnalyzer._check_header_footer_parts(
                (section.header, section.first_page_header, section.even_page_header),
                check_hyperlinks=True
            )
            have_links_in_header |= has_links
            
            # Check footers
            footer_info, _ = CVLayoutAnalyzer._check_header_footer_parts(
                (section.footer, section.first_page_footer, section.even_page_footer),
                check_hyperlinks=False
            )

        return HeaderFooterResult(
            has_header_info=header_info,
            has_footer_info=footer_info,
            have_links_in_header=have_links_in_header
        )

    @staticmethod
    def _check_header_footer_parts(
        parts: tuple, 
        check_hyperlinks: bool = False
    ) -> Tuple[bool, bool]:
        """Check header/footer parts for content and optionally hyperlinks.
        
        Args:
            parts: Tuple of header or footer objects
            check_hyperlinks: Whether to check for hyperlinks
        
        Returns:
            Tuple of (has_content, has_hyperlinks)
        """
        has_content = False
        has_hyperlinks = False
        
        for part in parts:
            try:
                if any(p.text.strip() for p in part.paragraphs):
                    has_content = True
                
                if check_hyperlinks:
                    for p in part.paragraphs:
                        if p.hyperlinks:
                            has_hyperlinks = True
            except Exception:
                pass
        
        return has_content, has_hyperlinks


    # ==================== Font & Style Methods ====================

    @staticmethod
    def get_effective_font_name(run: Run, paragraph, styles_element) -> Optional[str]:
        """Get the effective font name for a run, including inherited fonts.
        
        Follows python-docx style inheritance hierarchy:
        1. Direct run formatting (run.font.name)
        2. Paragraph style font
        3. Document default font from w:docDefaults
        
        References:
        - https://python-docx.readthedocs.io/en/latest/dev/analysis/features/text/font.html
        - https://github.com/python-openxml/python-docx/issues/383

        Args:
            run: The run object from python-docx
            paragraph: The paragraph object containing the run
            styles_element: The styles element from the document
        
        Returns:
            Font name string or None if not found
        """
        # 1. Check direct run formatting
        if run.font.name:
            return run.font.name

        # 2. Check paragraph style font
        if paragraph.style and hasattr(paragraph.style, 'font'):
            try:
                if paragraph.style.font.name:
                    return paragraph.style.font.name
            except:
                pass

        # 3. Check document defaults from XML
        try:
            rFonts_elements = styles_element.xpath(
                'w:docDefaults/w:rPrDefault/w:rPr/w:rFonts',
            )
            if rFonts_elements:
                rFonts = rFonts_elements[0]
                font_name = (
                    rFonts.get(f'{CVLayoutAnalyzer.WPML_NS}ascii') or
                    rFonts.get(f'{CVLayoutAnalyzer.WPML_NS}hAnsi') or
                    rFonts.get(f'{CVLayoutAnalyzer.WPML_NS}cs')
                )
                if font_name:
                    return font_name
        except Exception:
            pass

        return None

    @staticmethod
    def get_effective_font_size(run: Run) -> Optional[float]:
        """Get the effective font size for a run, including inherited sizes.
        
        Args:
            run: The run object from python-docx
        
        Returns:
            Font size in points or None if not found
        """
        if run.font and run.font.size:
            return run.font.size.pt

        style = run._parent.style
        while style:
            if style.font and style.font.size:
                return style.font.size.pt
            style = style.base_style

        return None

    # ==================== Text Content Analysis Methods ====================

    @staticmethod
    def _analyze_text_content(doc: Document) -> TextContentResult:
        """Analyze document text content for structure and formatting.
        
        Extracts:
        - All text with hyperlinks resolved
        - Font names and sizes
        - Tables, graphics detection
        - Header/footer presence
        
        Returns:
            TextContentResult dictionary
        """
        final_text: List[str] = []
        fonts: Set[str] = set()
        font_sizes: List[float] = []
        seen_uris: Set[str] = set()

        # Extract body text with hyperlinks
        CVLayoutAnalyzer._extract_text_recursive(
            doc._element, doc, final_text, seen_uris
        )

        # Extract fonts from paragraphs
        CVLayoutAnalyzer._extract_fonts_from_paragraphs(
            doc.paragraphs, doc.styles.element, fonts, font_sizes
        )

        # Extract fonts from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    CVLayoutAnalyzer._extract_fonts_from_paragraphs(
                        cell.paragraphs, doc.styles.element, fonts, font_sizes
                    )

        # Detect document features
        have_graphics = CVLayoutAnalyzer._detect_graphics(doc)
        have_images = CVLayoutAnalyzer._detect_images(doc)  
        have_textboxes = CVLayoutAnalyzer._detect_textboxes(doc)

        # Clean and join text
        extracted_text = CVLayoutAnalyzer._clean_text(final_text)

        return {
            "extracted_text": extracted_text,
            "fonts_names": list(fonts),
            "font_sizes": font_sizes,
            "have_hyperlinks": len(seen_uris) > 0,
        }

    @staticmethod
    def _get_text_from_element(elem) -> str:
        """Extract all text content from an XML element."""
        texts = [t.text for t in elem.iter() if t.tag.endswith('}t') and t.text]
        return " ".join(texts).strip()

    @staticmethod
    def _extract_text_recursive(
        element, 
        doc: Document, 
        final_text: List[str], 
        seen_uris: Set[str]
    ) -> None:
        """Recursively extract text from document elements (including tables).
        
        Args:
            element: XML element to extract from
            doc: Document object for relationship lookups
            final_text: List to accumulate text
            seen_uris: Set to track processed hyperlink URIs
        """
        for child in element.iterchildren():
            if isinstance(child, CT_P):
                para_text = CVLayoutAnalyzer._get_text_from_element(child)
                para_text = CVLayoutAnalyzer._resolve_hyperlinks(
                    child, doc, para_text, seen_uris
                )
                if para_text:
                    final_text.append(para_text)

            elif isinstance(child, CT_Tbl):
                # Extract text from table properly - iterate through paragraphs in table cells
                for para in child.findall(f".//{CVLayoutAnalyzer.WPML_NS}p"):
                    para_text = CVLayoutAnalyzer._get_text_from_element(para)
                    para_text = CVLayoutAnalyzer._resolve_hyperlinks(
                        para, doc, para_text, seen_uris
                    )
                    if para_text:
                        final_text.append(para_text)
            else:
                CVLayoutAnalyzer._extract_text_recursive(
                    child, doc, final_text, seen_uris
                )

    @staticmethod
    def _resolve_hyperlinks(
        para, 
        doc: Document, 
        para_text: str, 
        seen_uris: Set[str]
    ) -> str:
        """Resolve hyperlinks in paragraph text.
        
        Args:
            para: Paragraph XML element
            doc: Document object
            para_text: Original paragraph text
            seen_uris: Set to track processed URIs
        
        Returns:
            Text with hyperlink URLs appended to anchors
        """
        for hyperlink in para.findall(f".//{CVLayoutAnalyzer.WPML_NS}hyperlink"):
            rId = hyperlink.get(CVLayoutAnalyzer.RELATIONSHIPS_NS + "id")
            if rId and rId not in seen_uris:
                rel = doc.part.rels.get(rId)
                if rel and hasattr(rel, "target_ref"):
                    url = rel.target_ref
                    anchor = CVLayoutAnalyzer._get_text_from_element(hyperlink)
                    if anchor:
                        para_text = para_text.replace(anchor, f"{anchor} ({url})", 1)
                        seen_uris.add(rId)
        return para_text

    @staticmethod
    def _extract_fonts_from_paragraphs(
        paragraphs, 
        styles_element, 
        fonts: Set[str], 
        font_sizes: List[float]
    ) -> None:
        """Extract font information from paragraphs.
        
        Args:
            paragraphs: Iterable of paragraph objects
            styles_element: Document styles element
            fonts: Set to accumulate font names
            font_sizes: List to accumulate font sizes
        """
        for paragraph in paragraphs:
            for run in paragraph.runs:
                font_name = CVLayoutAnalyzer.get_effective_font_name(
                    run, paragraph, styles_element
                )
                if font_name:
                    fonts.add(font_name)

                font_size = CVLayoutAnalyzer.get_effective_font_size(run)
                if font_size and font_size > 0:
                    font_sizes.append(font_size)

    @staticmethod
    def _check_header_footer_content(doc: Document) -> Tuple[bool, bool]:
        """Check if document has content in headers or footers.
        
        Returns:
            Tuple of (has_header_info, has_footer_info)
        """
        has_header_info = False
        has_footer_info = False

        for section in doc.sections:
            header_parts = (section.header, section.first_page_header, section.even_page_header)
            footer_parts = (section.footer, section.first_page_footer, section.even_page_footer)
            
            has_header_info, _ = CVLayoutAnalyzer._check_header_footer_parts(header_parts)
            has_footer_info, _ = CVLayoutAnalyzer._check_header_footer_parts(footer_parts)

        return has_header_info, has_footer_info

    @staticmethod
    def _clean_text(text_lines: List[str]) -> str:
        """Clean and join text lines.
        
        Removes extra whitespace and joins non-empty lines.
        """
        cleaned = [
            re.sub(r'[ \t]+', ' ', line).strip()
            for line in text_lines
            if line.strip()
        ]

        cleaned = "\n".join(text_lines)

        replacements = {
            "â€“": "–",   # en dash
            "â€”": "—",   # em dash
            "â€˜": "'",   # left single quote
            "â€™": "'",   # right single quote
            "â€œ": '"',   # left double quote
            "â€": '"',   # right double quote
            "Â ": " ",    # non-breaking space artifact
            "Â": "",      # standalone stray marker
            "ï¼": ":",    # full-width colon corruption
            "ï½": ";",    # full-width semicolon corruption
        }

        for wrong, right in replacements.items():
            cleaned = cleaned.replace(wrong, right)

        BULLETS = r"[•●▪○◯◦◌⚫‣∙➢►▶◆◇■□✓✔➤]"
        cleaned = re.sub(BULLETS, "-", cleaned)
        # Capture the URL and reconstruct properly
        def fix_broken_urls(text):
            pattern = r'(\w+)?\s*([A-Za-z]+)?\s*\((https?://[^)]+)\)([A-Za-z0-9:/._-]+)?'

            def repl(m):
                prefix_word = m.group(1) or ""
                left_part   = m.group(2) or ""
                first_url   = m.group(3)
                right_part  = m.group(4) or ""

                # Case 1:
                # U (url)demy  -> Udemy (url)
                if right_part:
                    merged = left_part + right_part
                    return f"{prefix_word} {merged} ({first_url})"

                return m.group(0)

            return re.sub(pattern, repl, text)

        cleaned = fix_broken_urls(cleaned)
        cleaned = re.sub(r'[ \t]+', ' ', cleaned)  # Replace multiple spaces/tabs with single space
        cleaned = re.sub(r'[\uE000-\uF8FF]', ' ', cleaned) # Remove Unicode Private Use Area chars
        cleaned = cleaned.replace(" \n", "\n").replace("\n ", "\n")  # Clean spaces around newlines
        cleaned = re.sub(r'\n{2,}', '\n\n', cleaned)  # Limit consecutive newlines
        # cleaned = re.sub(r'\s+', ' ', cleaned)  # Remove extra whitespace
        cleaned = re.sub(r'\n\s*-\s*\n', '\n- ', cleaned) # fix standalone bullet lines
        cleaned = re.sub(r'\n(\S)\n', r' \1\n', cleaned) # If current line is a single character, append it to previous line
        return cleaned

    # ==================== Utility & Helper Methods ====================

    @staticmethod
    def _get_page_count_docx(doc: Document) -> int:
        """Estimate page count from DOCX by counting page breaks.
        
        Note: DOCX does not store page count explicitly, so this is an estimate.
        If the document contains no explicit page breaks, fall back to a
        layout-based estimation using `_estimate_pages_from_docx`.
        """
        page_count = 1
        try:
            page_count = sum(p.contains_page_break for p in doc.paragraphs) + 1
            est = CVLayoutAnalyzer._estimate_pages_from_docx(doc)
            return max(1, page_count, est)
        except Exception:
            return page_count

    @staticmethod
    def _estimate_pages_from_docx(doc: Document, page_height_in: float = 11.0) -> int:
        """Estimate number of pages for a DOCX by simulating layout height.

        This uses average section margins, average font size, paragraph
        spacing, table rows, and images to approximate how many pages the
        document would occupy for a given page height (in inches).
        """
        # SAFE SECTION HANDLING
        if doc.sections:
            top_margin = sum(s.top_margin.inches for s in doc.sections) / len(doc.sections)
            bottom_margin = sum(s.bottom_margin.inches for s in doc.sections) / len(doc.sections)
        else:
            top_margin = 1.0
            bottom_margin = 1.0

        usable_height = page_height_in - (top_margin + bottom_margin)
        if usable_height <= 0:
            return 1

        # FONT SIZE ESTIMATION
        font_sizes = []
        for p in doc.paragraphs:
            for r in p.runs:
                if r.font.size:
                    try:
                        font_sizes.append(r.font.size.pt)
                    except Exception:
                        pass

        avg_font = mean(font_sizes) if font_sizes else 12
        font_scale = (avg_font / 12)

        # HEIGHT ACCUMULATION MODEL
        total_height_units = 0.0

        for para in doc.paragraphs:
            text = (para.text or "").strip()
            if not text:
                continue

            word_count = len(text.split())

            # BASE LINE HEIGHT: assume ~0.2 inches per line at 12pt
            line_height = 0.2 * font_scale

            # estimate number of lines (rough average words per line)
            words_per_line = 10
            lines = max(1, word_count / words_per_line)

            para_height = lines * line_height

            # PARAGRAPH SPACING
            space_before = getattr(para.paragraph_format, 'space_before', None)
            space_after = getattr(para.paragraph_format, 'space_after', None)
            sb = getattr(space_before, 'pt', 0) if space_before is not None else 0
            sa = getattr(space_after, 'pt', 0) if space_after is not None else 0
            spacing_height = ((sb + sa) / 72)

            # HEADINGS (more spacing)
            style_name = (para.style.name or "").lower() if para.style else ""
            if "heading" in style_name:
                para_height *= 1.6
                spacing_height += 0.15

            # LISTS
            try:
                if para._p.xpath(".//w:numPr"):
                    para_height *= 1.2
            except Exception:
                pass

            total_height_units += para_height + spacing_height

        # TABLES (dense but tall blocks)
        for table in doc.tables:
            rows = len(table.rows)
            table_height = rows * 0.35  # inches per row
            total_height_units += table_height

        # IMAGES (major page impact)
        image_count = 0
        try:
            for rel in doc.part.rels.values():
                if hasattr(rel, "target_ref") and "image" in rel.target_ref:
                    image_count += 1
        except Exception:
            pass

        total_height_units += image_count * 1.2

        # FINAL PAGE ESTIMATION
        estimated_pages = math.ceil(total_height_units / usable_height)
        return max(1, int(estimated_pages))

    @staticmethod
    def _detect_section_breaks(doc: Document) -> Optional[List[int]]:
        """Detect section breaks in the document to determine page boundaries.
        
        Returns:
            List of page numbers where each section breaks, or None if unable to detect
        """
        try:
            section_breaks = []
            current_page = 1
            
            for paragraph in doc.paragraphs:
                pPr = paragraph._element.find(f"{CVLayoutAnalyzer.WPML_NS}pPr")
                if pPr is not None:
                    sectPr = pPr.find(f"{CVLayoutAnalyzer.WPML_NS}sectPr")
                    if sectPr is not None:
                        section_breaks.append(current_page)
            
            return section_breaks if section_breaks else None
        except Exception:
            return None
    
    @staticmethod
    def _get_section_index_for_page(page_num: int, section_breaks: List[int]) -> int:
        """Determine which section a page belongs to based on section breaks.
        
        Args:
            page_num: The page number (1-indexed)
            section_breaks: List of page numbers where sections start/break
        
        Returns:
            Index of the section this page belongs to
        """
        section_idx = 0
        for break_page in sorted(section_breaks):
            if page_num >= break_page:
                section_idx += 1
            else:
                break
        return section_idx
    
    @staticmethod
    def word_count(text: str) -> int:
        """Count words in text string.
        
        Args:
            text: Input text
        
        Returns:
            Number of words
        """
        if not text:
            return 0
        return len(text.split())